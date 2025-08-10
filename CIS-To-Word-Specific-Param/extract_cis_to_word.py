#!/usr/bin/env python3
"""
CIS Benchmark PDF → Word (Hierarchical + Intros, Robust + Borders v6)

Fixes in v6
-----------
• **No missing controls**: more robust control detection
  - Detects (Automated|Manual) even when it wraps to next lines
  - Only treats a numeric line as a control if its own header (or the
    immediate follow-up header lines) include the status OR if a
    "Description:" appears **before the next numeric header**
  - Prevents misclassifying structural headers (e.g., 4.1.2) as controls
• **Borders that actually show**: each table uses the built‑in "Table Grid"
  style **and** explicit XML borders (top/left/bottom/right/insideH/insideV)
• **Accurate headers + intro text**: header text = header line only; the
  narrative sentence(s) directly after a structural header are kept as
  introductory paragraphs above the table for that header
• **Strict numeric ordering** and deepest‑parent grouping for hierarchy

Usage
-----
python cis_pdf_word_hier_robust_v6.py \
  --pdf "/mnt/data/CIS_IBM_WebSphere_Liberty_Benchmark_v1.0.0.pdf" \
  --out "CIS_WebSphere_Hier_Robust_v6.docx"

Requires: pip install pdfminer.six python-docx
"""
from __future__ import annotations

import argparse
import re
import warnings
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple

warnings.filterwarnings(
    "ignore",
    message=r"CropBox missing from /Page, defaulting to MediaBox",
)

try:
    from pdfminer.high_level import extract_text
except Exception as e:
    raise SystemExit("pdfminer.six is required. Install it with: pip install pdfminer.six\n" + str(e))

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except Exception as e:
    raise SystemExit("python-docx is required. Install it with: pip install python-docx\n" + str(e))

# ------------------ Patterns ------------------
NUMERIC_HEADER_RX   = re.compile(r"^(?P<code>\d+(?:\.\d+)*)\s+(?P<title>.*)$")
# Strict control signature if all on one line
CONTROL_HEADER_RX   = re.compile(r"^(?P<code>\d+(?:\.\d+)*)\s+(?P<title>.+?)\s*\((?:Automated|Manual)\)\s*$")
SECTION_NAMES       = [
    "Description","Rationale","Impact","Audit","Default Value",
    "Remediation","References","Profile Applicability","CIS Controls"
]
TOC_TITLE_RX        = re.compile(r"^Table of Contents$", re.IGNORECASE)
TOC_DOT_LEADER_RX   = re.compile(r"\.\.\.+\s*\d+$")
TRAILING_PAGE_NUM_RX= re.compile(r"\s+\d+$")
PAGE_FOOTER_RX      = re.compile(r"^Page\s+\d+\s*$", re.IGNORECASE)

TABLE_HEADER_STRINGS = {
    "control reference id",
    "control name & description",
    "description",
    "remediation",
}

# ------------------ Helpers ------------------
def collapse_spaces(s: str) -> str:
    return re.sub(r"[ \t]+", " ", s).strip()


def normalize_text(text: str) -> str:
    text = text.replace("\u00A0", " ")
    text = text.replace("\r", "")
    # de-hyphenate wraps like "confi-\nguration"
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    return text


def is_section_header(line: str) -> bool:
    s = line.strip().rstrip(":")
    return any(s.lower() == name.lower() for name in SECTION_NAMES)


def is_toc_like(line: str) -> bool:
    if TOC_DOT_LEADER_RX.search(line):
        return True
    if PAGE_FOOTER_RX.match(line):
        return True
    if TRAILING_PAGE_NUM_RX.search(line) and NUMERIC_HEADER_RX.match(line):
        return True
    return False


def is_table_header_line(line: str) -> bool:
    return line.strip().lower() in TABLE_HEADER_STRINGS


def is_url_footnote(line: str) -> bool:
    m = NUMERIC_HEADER_RX.match(line)
    if not m:
        return False
    title = m.group("title").strip().lstrip(".").lstrip()
    return title.lower().startswith("http://") or title.lower().startswith("https://")


def text_to_lines(text: str) -> List[str]:
    lines: List[str] = []
    buf = ""
    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            if buf:
                lines.append(collapse_spaces(buf))
                buf = ""
            continue
        if (
            NUMERIC_HEADER_RX.match(line) or is_section_header(line) or
            TOC_TITLE_RX.match(line.strip()) or is_toc_like(line) or
            PAGE_FOOTER_RX.match(line) or is_url_footnote(line) or is_table_header_line(line)
        ):
            if buf:
                lines.append(collapse_spaces(buf))
                buf = ""
            # Drop URL footnotes and static table headers and "Table of Contents" title
            if not (is_url_footnote(line) or is_table_header_line(line) or TOC_TITLE_RX.match(line)):
                lines.append(collapse_spaces(line))
            continue
        if buf:
            sep = "" if buf.endswith(" ") else " "
            buf += sep + line.strip()
        else:
            buf = line.strip()
    if buf:
        lines.append(collapse_spaces(buf))
    # final scrub
    return [ln for ln in lines if not (TOC_TITLE_RX.match(ln) or is_toc_like(ln) or is_table_header_line(ln))]


def dotted_key(code: str) -> Tuple[int, ...]:
    return tuple(int(x) for x in code.split('.'))

# ------------------ Data classes ------------------
@dataclass
class Anchor:
    code: str
    title: str  # from SAME line only
    start: int
    end: int

@dataclass
class Control:
    cid: str
    title: str
    start: int
    end: int

@dataclass
class StructHeader:
    code: str
    title: str
    start: int
    end: int

# ------------------ Parsing ------------------

def gather_anchors(lines: List[str]) -> List[Anchor]:
    raw_idxs = [i for i, ln in enumerate(lines) if NUMERIC_HEADER_RX.match(ln) and not is_toc_like(ln) and not is_url_footnote(ln)]
    # Keep first occurrence per code
    seen: Dict[str, int] = {}
    idxs: List[int] = []
    for i in raw_idxs:
        code = NUMERIC_HEADER_RX.match(lines[i]).group("code")
        if code not in seen:
            seen[code] = i
            idxs.append(i)
    idxs.sort()

    anchors: List[Anchor] = []
    for p, i in enumerate(idxs):
        m = NUMERIC_HEADER_RX.match(lines[i])
        code  = m.group("code")
        title = m.group("title").strip()
        end   = idxs[p + 1] if p + 1 < len(idxs) else len(lines)
        anchors.append(Anchor(code=code, title=title, start=i, end=end))
    return anchors


def header_block_text(a: Anchor, lines: List[str], span: int = 3) -> str:
    """Return the header line plus up to `span` following non-structural lines
    (stops early on next numeric or a section header)."""
    parts = [lines[a.start]]
    j = a.start + 1
    taken = 0
    while j < a.end and taken < span:
        ln = lines[j]
        if NUMERIC_HEADER_RX.match(ln) or is_section_header(ln) or is_toc_like(ln) or is_table_header_line(ln) or is_url_footnote(ln):
            break
        parts.append(ln)
        taken += 1
        j += 1
    return collapse_spaces(" ".join(parts))


def is_control_anchor(a: Anchor, lines: List[str]) -> bool:
    # 1) Strong hint: header (with up to next 3 lines) contains (Automated|Manual)
    hb = header_block_text(a, lines, span=3)
    if re.search(r"\((?:Automated|Manual)\)\s*$", hb):
        return True
    # 2) Fallback: within the block AFTER the header but BEFORE the next numeric header
    #    we encounter a "Description:" section — that indicates a control body
    j = a.start + 1
    while j < a.end:
        ln = lines[j]
        if NUMERIC_HEADER_RX.match(ln):  # stop at the very next numeric header
            break
        if is_section_header(ln) and ln.strip().rstrip(":").lower() == "description":
            return True
        j += 1
    return False


def classify(anchors: List[Anchor], lines: List[str]) -> Tuple[List[Control], List[StructHeader]]:
    controls: Dict[str, Control] = {}
    structs : Dict[str, StructHeader] = {}
    for a in anchors:
        if is_control_anchor(a, lines):
            # Title from header block (includes wrapped title if it spilled)
            hb = header_block_text(a, lines, span=3)
            # Strip leading code
            t = NUMERIC_HEADER_RX.match(hb).group("title").strip()
            controls.setdefault(a.code, Control(cid=a.code, title=t, start=a.start, end=a.end))
        else:
            structs.setdefault(a.code, StructHeader(code=a.code, title=a.title, start=a.start, end=a.end))
    return list(controls.values()), list(structs.values())


def extract_intro_for_struct(s: StructHeader, lines: List[str]) -> str:
    intro: List[str] = []
    j = s.start + 1
    while j < s.end:
        ln = lines[j]
        if NUMERIC_HEADER_RX.match(ln) or is_section_header(ln) or is_toc_like(ln) or is_table_header_line(ln) or is_url_footnote(ln):
            break
        intro.append(ln)
        j += 1
    return "\n".join(intro).strip()


def extract_section(block: List[str], section_name: str) -> str:
    try:
        idx = next(i for i, ln in enumerate(block) if ln.strip().rstrip(":").lower() == section_name.lower())
    except StopIteration:
        return ""
    body: List[str] = []
    j = idx + 1
    while j < len(block):
        if is_section_header(block[j]) or NUMERIC_HEADER_RX.match(block[j]) or is_toc_like(block[j]) or is_table_header_line(block[j]) or is_url_footnote(block[j]):
            break
        body.append(block[j])
        j += 1
    return "\n".join(body).strip()


def longest_structural_parent(code: str, struct_map: Dict[str, StructHeader]) -> Optional[str]:
    parts = code.split('.')
    for L in range(len(parts) - 1, 0, -1):
        pref = '.'.join(parts[:L])
        if pref in struct_map:
            return pref
    if parts[0] in struct_map:
        return parts[0]
    return None

# ------------------ Borders ------------------

def set_table_borders(table) -> None:
    """Apply both style and explicit XML borders so borders show everywhere."""
    # Style-based borders
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    # Explicit XML borders on the table
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '8')   # 8 = ~0.5pt
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '000000')
        tblBorders.append(elem)
    # Remove existing borders if any, then append
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
    tblPr.append(tblBorders)

# ------------------ Word export ------------------

def write_word(controls: List[Control], structs: List[StructHeader], lines: List[str], out_path: Path) -> None:
    struct_map: Dict[str, StructHeader] = {s.code: s for s in structs}
    # Ensure we don't drop controls that somehow lack a detected parent
    buckets: Dict[str, List[Control]] = {code: [] for code in struct_map.keys()}
    orphans: List[Control] = []

    for c in controls:
        parent = longest_structural_parent(c.cid, struct_map)
        if parent is not None:
            buckets[parent].append(c)
        else:
            orphans.append(c)

    ordered_structs = sorted(structs, key=lambda s: dotted_key(s.code))

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.font.size = Pt(11)

    for s in ordered_structs:
        controls_here = sorted(buckets.get(s.code, []), key=lambda c: dotted_key(c.cid))
        intro = extract_intro_for_struct(s, lines)
        if not controls_here and not intro:
            continue
        depth = min(len(s.code.split('.')), 5)
        doc.add_heading(f"{s.code} {s.title}", level=depth)
        if intro:
            doc.add_paragraph(intro)
        if controls_here:
            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = 'Control Reference ID'
            hdr[1].text = 'Control Name & Description'
            hdr[2].text = 'Description'
            hdr[3].text = 'Remediation'
            for c in controls_here:
                block = lines[c.start:c.end]
                desc = extract_section(block, "Description")
                rem  = extract_section(block, "Remediation")
                row = table.add_row().cells
                row[0].text = c.cid
                row[1].text = c.title
                row[2].text = desc
                row[3].text = rem
            set_table_borders(table)
            doc.add_paragraph("")

    # If any orphan controls remain, dump them at the end under a catch-all
    if orphans:
        doc.add_heading("Unplaced Controls", level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = 'Control Reference ID'
        hdr[1].text = 'Control Name & Description'
        hdr[2].text = 'Description'
        hdr[3].text = 'Remediation'
        for c in sorted(orphans, key=lambda c: dotted_key(c.cid)):
            block = lines[c.start:c.end]
            desc = extract_section(block, "Description")
            rem  = extract_section(block, "Remediation")
            row = table.add_row().cells
            row[0].text = c.cid
            row[1].text = c.title
            row[2].text = desc
            row[3].text = rem
        set_table_borders(table)

    doc.save(str(out_path))

# ------------------ CLI ------------------

def main():
    ap = argparse.ArgumentParser(description="Export CIS controls to hierarchical Word with robust detection + borders")
    ap.add_argument("--pdf", required=True, help="Path to CIS Benchmark PDF")
    ap.add_argument("--out", required=True, help="Output .docx path")
    args = ap.parse_args()

    pdf_path = Path(args.pdf)
    if not pdf_path.exists():
        raise SystemExit(f"PDF not found: {pdf_path}")

    raw = extract_text(str(pdf_path))
    raw = normalize_text(raw)
    lines = text_to_lines(raw)

    anchors = gather_anchors(lines)
    controls, structs = classify(anchors, lines)

    if not controls:
        raise SystemExit("No controls detected. Try adjusting detection heuristics.")

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    write_word(controls, structs, lines, out_path)
    print(f"Saved hierarchical Word with borders to {args.out}")

if __name__ == "__main__":
    main()
